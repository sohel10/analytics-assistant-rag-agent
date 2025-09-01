import os, json, re
from io import BytesIO
from pathlib import Path
from typing import List, Dict, Any

import streamlit as st
import pandas as pd
from docx import Document

import torch
from transformers import pipeline

# Quieter HF symlink warning on Linux
os.environ.setdefault("HF_HUB_DISABLE_SYMLINKS_WARNING", "1")

# -----------------------------
# Seed reference snippets (RAG-lite)
# -----------------------------
SOURCES_DIR = Path("sources")
SOURCES_DIR.mkdir(parents=True, exist_ok=True)

DEFAULT_SOURCES = {
    "visit_cadence.txt": """# sources/visit_cadence.txt
TITLE: Visit Cadence & Windows (Illustrative)
PURPOSE - Define standard visit schedule, allowable windows, and required assessments.
STANDARD SCHEDULE & WINDOWS
- Screening: Day −28 to Day −1 (window ±0)
- Baseline/Day 1: Day 1 (±0)
- Week 2: Day 14 (±3 days)
- Week 4: Day 28 (±7 days)
- Then Every 4 Weeks: q4w from Day 28 (±7 days)
- End of Treatment (EOT): within 7 days after last dose
- Safety Follow-up: 30 (±7) days after EOT
REQUIRED PROCEDURES (BY VISIT TYPE)
- All On-Treatment Visits (incl. Week 2, Week 4, q4w):
  • Vitals, weight, ECOG/PS
  • ConMeds update
  • AE/SAE review since prior visit
  • IP accountability & dosing review
  • Labs: CBC/CMP (or per protocol)
MISSED / OUT-OF-WINDOW HANDLING
- Schedule earliest feasible date; document deviation & rationale in source/EDC.
TELEVISIT ALLOWANCE
- Televisit may be used for wellness/AE checks when appropriate.
KEYWORDS visit window, schedule, q4w, EOT, follow-up, deviation, televisit
""",
    "safety_reporting.txt": """# sources/safety_reporting.txt
TITLE: Safety Reporting (AE/SAE/SUSAR) – Illustrative Guidance
REPORTING
- All AEs: document each visit (onset, severity, outcome, relatedness).
- SAEs: report to Sponsor Safety within 24 hours of awareness.
- SUSAR: expedited timelines (7 or 15 calendar days) per seriousness.
KEYWORDS AE, SAE, SUSAR, CTCAE v5, expedited, 24 hours, 7-day, 15-day
""",
    "followup_cadence.txt": "Routine follow-up every 3 months or per PI discretion. Encourage BP logs for hypertension.\n",
    "diabetes_basics.txt": """TITLE: Diabetes Care (Illustrative)
- A1c goal individualized; monitor per clinician.
- Annual eye & foot exams; daily foot checks.
- Hypoglycemia education if on insulin/sulfonylureas.
- Lifestyle: balanced diet, regular activity, weight management.
KEYWORDS diabetes, A1c, hypoglycemia, foot exam, eye exam
"""
}

def ensure_seed_sources():
    for fname, text in DEFAULT_SOURCES.items():
        f = SOURCES_DIR / fname
        if not f.exists():
            f.write_text(text, encoding="utf-8")

ensure_seed_sources()

# -----------------------------
# Tiny retriever (token overlap)
# -----------------------------
def load_sources() -> Dict[str, str]:
    return {p.name: p.read_text(encoding="utf-8")
            for p in SOURCES_DIR.glob("*.txt")}

TOKEN_RX = re.compile(r"[A-Za-z0-9\-\+%]+")

def tokens(s: str) -> set:
    return {t.lower() for t in TOKEN_RX.findall(s or "")}

def build_query(patient: Dict[str, Any]) -> str:
    parts = []
    demo = patient.get("demographics", {})
    if "age_years" in demo: parts += ["age", str(demo["age_years"])]
    if "gender" in demo: parts += [str(demo["gender"])]
    conds = [c.get("description","") for c in patient.get("conditions", []) if c.get("active")]
    if conds:
        parts += ["conditions"] + conds
    parts += ["eligibility", "visit cadence", "safety reporting"]
    return " ".join(parts)

def retrieve_top_k(query: str, corpus: Dict[str,str], k: int = 3):
    q = tokens(query)
    scored = []
    for name, text in corpus.items():
        score = len(q & tokens(text))
        scored.append((name, score))
    scored.sort(key=lambda x: x[1], reverse=True)
    return [(n, s, corpus[n]) for n, s in scored[:k]]

# -----------------------------
# LLM (FLAN-T5)
# -----------------------------
@st.cache_resource(show_spinner=True)
def load_pipe(model_name: str = "google/flan-t5-base"):
    device = 0 if torch.cuda.is_available() else -1
    return pipeline("text2text-generation", model=model_name, device=device)

def build_prompt(patient: Dict[str, Any], citations: List[str]) -> str:
    facts = {
        "patient_id": patient.get("patient_id"),
        "demographics": patient.get("demographics", {}),
        "conditions": (patient.get("conditions") or [])[:5],
    }
    cites = "\n".join(f"- {c}" for c in citations) if citations else "- (none)"
    return f"""
You are a medical assistant writing a brief, friendly note to a patient.
Use ONLY the facts supplied plus the cited reference snippets. Do not invent data.
Write 4–6 sentences, plain language. End with: "Contact your care team if symptoms change."

Facts (JSON):
{json.dumps(facts, indent=2)}

Reference snippets (titles only):
{cites}

Now write the patient message (no headings, no bullets):
""".strip()

def generate_message(pipe, patient_json: str, top_k: int = 3) -> Dict[str, Any]:
    try:
        patient = json.loads(patient_json)
    except Exception as e:
        return {"error": f"Invalid JSON: {e}"}

    query = build_query(patient)
    corpus = load_sources()
    results = retrieve_top_k(query, corpus, k=top_k)
    cited_titles = [name for name, _, _ in results]

    prompt = build_prompt(patient, cited_titles)
    out = pipe(
        prompt,
        max_new_tokens=200,
        no_repeat_ngram_size=3,
        length_penalty=0.9,
        num_beams=4,
        do_sample=False,
    )
    text = out[0]["generated_text"].strip()
    if not text.endswith("Contact your care team if symptoms change."):
        if not text.endswith("."): text += "."
        text += " Contact your care team if symptoms change."

    return {
        "query": query,
        "snippets": [{"name": n, "score": s, "text": t} for n, s, t in results],
        "message": text,
    }

# -----------------------------
# DOCX export
# -----------------------------
def to_docx(title: str, text: str) -> bytes:
    doc = Document()
    doc.add_heading(title, level=1)
    for para in text.split("\n"):
        para = para.strip()
        if para:
            doc.add_paragraph(para)
    buf = BytesIO(); doc.save(buf); buf.seek(0)
    return buf.read()

# -----------------------------
# Streamlit UI
# -----------------------------
st.set_page_config(page_title="RAG-lite Clinical Assistant", page_icon="🩺", layout="centered")
st.title("🩺 RAG-lite Clinical Assistant (FLAN-T5)")

with st.sidebar:
    st.subheader("Model")
    model_choice = st.selectbox(
        "FLAN-T5 model",
        ["google/flan-t5-base", "google/flan-t5-small", "google/flan-t5-large"],
        index=0,
        help="Cloud will use CPU; local CUDA will be used if available."
    )
    pipe = load_pipe(model_choice)
    st.caption("Model loaded.")
    top_k = st.slider("Snippets to cite", 1, 5, 3, 1)

tabs = st.tabs(["Patient narrative", "Batch", "Manage sources"])

# ---- Patient narrative ----
with tabs[0]:
    example = {
        "patient_id": "X001",
        "demographics": {"age_years": 55, "gender": "M"},
        "conditions": [{"description": "Diabetes", "active": True}],
    }
    st.markdown("Paste a patient **JSON**:")
    txt = st.text_area("Patient JSON", value=json.dumps(example, indent=2), height=220)
    if st.button("Generate message"):
        res = generate_message(pipe, txt, top_k=top_k)
        if "error" in res:
            st.error(res["error"])
        else:
            st.markdown("**Query terms**")
            st.code(res["query"])
            st.markdown("**Top snippets**")
            for sn in res["snippets"]:
                with st.expander(f"{sn['name']} (score={sn['score']})"):
                    st.write(sn["text"])
            st.markdown("**Generated message**")
            st.write(res["message"])
            try:
                pid = json.loads(txt).get("patient_id","")
            except Exception:
                pid = ""
            st.download_button(
                "Download DOCX",
                to_docx(f"Patient {pid}", res["message"]),
                file_name=f"patient_{pid or 'message'}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

# ---- Batch (CSV) ----
with tabs[1]:
    st.markdown("Upload a CSV with at least a **patient_json** column.")
    up = st.file_uploader("Upload CSV", type=["csv"])
    if up is not None:
        df = pd.read_csv(up)
        if "patient_json" not in df.columns:
            st.error("CSV must include a 'patient_json' column.")
        else:
            out = []
            for _, row in df.iterrows():
                js = str(row["patient_json"])
                res = generate_message(pipe, js, top_k=top_k)
                msg = res.get("message","") if isinstance(res, dict) else str(res)
                try:
                    pid = json.loads(js).get("patient_id","")
                except Exception:
                    pid = ""
                out.append({"patient_id": pid, "message": msg})
            out_df = pd.DataFrame(out)
            st.dataframe(out_df, use_container_width=True)
            st.download_button(
                "Download results CSV",
                out_df.to_csv(index=False).encode("utf-8"),
                file_name="raglite_batch_results.csv",
                mime="text/csv"
            )

# ---- Manage sources ----
with tabs[2]:
    st.markdown("Reference files in **./sources** used for retrieval:")
    corpus = load_sources()
    if corpus:
        for name, text in corpus.items():
            with st.expander(name):
                st.code(text)
    else:
        st.info("No .txt files found in ./sources")

    st.markdown("---")
    st.markdown("Add a new source (.txt):")
    uploaded = st.file_uploader("Upload .txt", type=["txt"], key="src_up")
    if uploaded is not None:
        dest = SOURCES_DIR / uploaded.name
        dest.write_bytes(uploaded.getbuffer())
        st.success(f"Saved: sources/{uploaded.name}. Press Rerun to refresh the list.")
